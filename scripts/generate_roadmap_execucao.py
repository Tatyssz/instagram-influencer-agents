#!/usr/bin/env python3
"""Gera ROADMAP-EXECUCAO.md e ROADMAP-EXECUCAO.docx a partir de roadmap_execucao_data.py."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from roadmap_execucao_data import (  # noqa: E402
    AUTHOR,
    DOC_VERSION,
    EPICS,
    PREREQUISITES_DONE,
    STATUS_LABEL,
    VERSION_DATE,
    Epic,
    Task,
)

OUT_DIR = ROOT / "docs" / "saas"
MD_PATH = OUT_DIR / "ROADMAP-EXECUCAO.md"
DOCX_PATH = OUT_DIR / "ROADMAP-EXECUCAO.docx"


def _criteria_md(task: Task) -> str:
    lines = []
    for checked, text in zip(task.criteria_checked, task.acceptance_criteria, strict=True):
        mark = "x" if checked else " "
        lines.append(f"  - [{mark}] {text}")
    return "\n".join(lines)


def _criteria_docx_symbol(checked: bool) -> str:
    return "\u2611" if checked else "\u2610"  # ☑ / ☐


def task_md(task: Task) -> str:
    deps = ", ".join(task.dependencies) if task.dependencies else "Nenhuma"
    status = STATUS_LABEL.get(task.status, task.status)
    return f"""### [{task.id}] {task.title}

| Campo | Conteúdo |
|-------|----------|
| **Status** | {status} |
| **Dependências** | {deps} |

**Descrição:** {task.description}

**Resultado esperado:** {task.expected_result}

**Critérios de aceite:**

{_criteria_md(task)}

---
"""


def epic_done_count(epic: Epic) -> int:
    return sum(1 for t in epic.tasks if t.status == "done")


def epic_md(epic: Epic) -> str:
    done = epic_done_count(epic)
    total = len(epic.tasks)
    icon = "✅" if done == total else "🔄" if done else "⬜"
    lines = [
        f"## {epic.id} — {epic.title}",
        "",
        f"**Objetivo:** {epic.goal}  ",
        f"**Duração:** {epic.duration}  ",
        f"**Progresso épico:** {icon} {done}/{total}",
        "",
    ]
    for t in epic.tasks:
        lines.append(task_md(t))
    return "\n".join(lines)


def prerequisites_md() -> str:
    lines = [
        "## Pré-requisitos concluídos (toolkit @tatyzacharias)",
        "",
        "*Motor do SaaS — feito antes deste roadmap.*",
        "",
    ]
    for item in PREREQUISITES_DONE:
        lines.append(f"- [x] {item}")
    lines.append("")
    return "\n".join(lines)


def generate_markdown() -> str:
    header = f"""# Roadmap de execução — ComJuntas SaaS

> **Autora:** {AUTHOR}  
> **Versão:** {DOC_VERSION}  
> **Data {DOC_VERSION}:** {VERSION_DATE}  
> **Formato:** Épico (bloco) = história Jira · Tarefa = card  
> **Fonte:** `scripts/roadmap_execucao_data.py` → regenere com `python scripts/generate_roadmap_execucao.py`  
> **Índice geral:** [docs/README.md](../README.md)

---

## Como usar

1. Marque `[x]` nos critérios de aceite conforme concluir.
2. **Status** do card: ⬜ A fazer · 🔄 Em progresso · ✅ Concluído
3. Não pule épico sem critério de saída do anterior (F0→F1→…).

---

"""
    body = "\n".join(epic_md(e) for e in EPICS)
    prereq = prerequisites_md()
    footer = """
## Legenda de dependências

- IDs referem-se a cards (ex.: `F0-08`, `DEC-01`)
- Épico paralelo **EPIC-P** não bloqueia SaaS; **TK-05** depende de **F2-01**

"""
    return header + prereq + body + footer


def generate_docx() -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Roadmap de Execução — ComJuntas SaaS", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    meta = doc.add_paragraph()
    meta.add_run(f"Autora: {AUTHOR}\n").bold = True
    meta.add_run(f"Versão: {DOC_VERSION}\n")
    meta.add_run(f"Data {DOC_VERSION}: {VERSION_DATE}\n")
    meta.add_run(
        "Checklist executável · Épico = História Jira · Tarefa = Card\n"
        "Marque ☐ → ☑ conforme conclusão"
    )

    doc.add_heading("Pré-requisitos concluídos (toolkit)", level=1)
    doc.add_paragraph("Motor do SaaS — feito antes deste roadmap.")
    for item in PREREQUISITES_DONE:
        doc.add_paragraph(f"{_criteria_docx_symbol(True)} {item}")

    for epic in EPICS:
        doc.add_page_break()
        done = epic_done_count(epic)
        total = len(epic.tasks)
        doc.add_heading(f"{epic.id} — {epic.title}", level=1)
        doc.add_paragraph(f"Objetivo: {epic.goal}")
        doc.add_paragraph(f"Duração: {epic.duration}")
        p = doc.add_paragraph("Progresso épico: ")
        p.add_run(f"{done}/{total} concluídos").bold = True

        for task in epic.tasks:
            doc.add_heading(f"[{task.id}] {task.title}", level=2)

            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Status"
            table.cell(0, 1).text = STATUS_LABEL.get(task.status, task.status)
            table.cell(1, 0).text = "Dependências"
            table.cell(1, 1).text = (
                ", ".join(task.dependencies) if task.dependencies else "Nenhuma"
            )

            doc.add_paragraph("Descrição:", style="Heading 3")
            doc.add_paragraph(task.description)

            doc.add_paragraph("Resultado esperado:", style="Heading 3")
            doc.add_paragraph(task.expected_result)

            doc.add_paragraph("Critérios de aceite:", style="Heading 3")
            for checked, text in zip(
                task.criteria_checked, task.acceptance_criteria, strict=True
            ):
                doc.add_paragraph(f"{_criteria_docx_symbol(checked)} {text}")

            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(generate_markdown(), encoding="utf-8")
    print(f"MD  -> {MD_PATH}")
    generate_docx()
    print(f"DOCX -> {DOCX_PATH}")


if __name__ == "__main__":
    main()
